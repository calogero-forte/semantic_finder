"""
License Statement & Module Information
======================================

This code is provided as open-source software and has been developed as part of the 
Master in Applied Artificial Intelligence postgraduate course, for the Python Programming topic.

The purpose of this application is to serve as a Model Context Protocol (MCP) server, 
providing a Large Language Model (LLM) the capability to access and retrieve 
information from local documents to answer related queries.

- Program Name: Semantic Finder
- Module Name: word_document.py
- Revision: 1.0
- Author: Calogero Forte
- Affiliation: University of Palermo
- Development Date: May 2026
"""

from docx.opc.constants import RELATIONSHIP_TYPE
import xml.etree.ElementTree as ET
from .document import Document, DocumentException
from .toc_document import TocDocument
from .utils import *
import os
import re
import math
import logging
import docx

logger = logging.getLogger(__name__)

class WordDocument(Document, TocDocument):
    """
    This class maintains an internal representation 
    of a DOCX file and provides methods to search 
    through it and extract text per pages and sections.
    """

    def __init__(self, docx_path_i):
        """
        Initialize a WordDocument

        docx_path_i: (str) The path (absolute or relative) 
                     of the docx file to open

        Return
        -------------------
        A WordDocument
        """
        super().__init__(path_i=docx_path_i)

        if not os.path.exists(docx_path_i):
            raise DocumentException(f"File not found: {docx_path_i}")

        if not docx_path_i.lower().endswith('.docx'):
            raise DocumentException("The given file is not a DOCX")

        try:
            self.doc = docx.Document(docx_path_i)
            self.__toc = self.__build_toc()

        except Exception as e:
            raise DocumentException(f"Failed to load DOCX: {e}")

        

    ##########################################################
    # Public methods
    ##########################################################

    # Redefining parent methods

    @property
    def page_num(self) -> int:
        page_count = 1
        try:
            part = self.doc.part.package.part_related_by(RELATIONSHIP_TYPE.EXTENDED_PROPERTIES)
            xml_str = part.blob.decode('utf-8')
            root = ET.fromstring(xml_str)
            ns = {'ep': 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'}
            pages_elem = root.find('ep:Pages', ns)
            if pages_elem is not None and pages_elem.text.isdigit():
                page_count = max(1, int(pages_elem.text))
                return page_count
        except Exception as e:
            logger.warning(f"Could not retrieve page count from metadata, defaulting to 1. Error: {e}")

    ##################################################

    @page_num.setter
    def page_num(self, value_i: int) -> None:
        raise Exception("The page number cannot be modified")

    ##################################################

    @property
    def toc(self) -> list:
        if self.__toc is None or len(self.__toc) <= 1: # Only EOF is present
            raise DocumentException("The table of contents is empty.")
        return self.__toc

    ##################################################

    def get_section_text(self, section_frame_i: dict, fine_trimming_i: bool = False) -> str:
        """
        Get the text content of a pages frame

        section_frame_i: (dict) the output of search_by_title(.)

        fine_trimming_i: (bool) if true, the method calls private sub-rutines 
                         to better clean the extracted text. Default is False.

        Return
        -------------------
        pages_text_o: (str) the text content of the pages frame.
                      It may be correctly empty.
        """
        page_text_o = ""
        pages_text_list = []

        # Section extremes
        start_page = section_frame_i['start']['page']
        start_index = section_frame_i['start']['index']
        end_page = section_frame_i['end']['page']
        end_index = section_frame_i['end']['index']

        # Strings to compare
        sec_title_plain = str( self.toc[start_index][1] )
        sec_title_norm = normalize_string( sec_title_plain )
        next_sec_title_norm = normalize_string( str(self.toc[end_index][1] ) )

        extracting = False

        for par in self.doc.paragraphs:
            
            try:

                curr_lines = []
                
                if ( par.style.name.startswith('Heading') and 
                    self.__extract_heading_level(par.style.name) == self.toc[start_index][0] and 
                    sec_title_norm in normalize_string(par.text) ):
                    
                    extracting = True
                    continue

                if ( par.style.name.startswith('Heading') and 
                    self.__extract_heading_level(par.style.name) == self.toc[end_index][0] and
                    next_sec_title_norm in normalize_string(par.text) ):

                    extracting = False
                    break

                if extracting:
                    curr_lines = par.text.splitlines()

                    if fine_trimming_i:
                        curr_lines = remove_short_lines(curr_lines)

            except (ValueError, IndexError) as e:
                logger.error(f"Error trimming text from frame [{start_page}, {end_page}]: {e}")
                curr_lines = []
            except Exception as e:
                logger.error(f"Error extracting text from frame [{start_page}, {end_page}]: {e}")
                curr_lines = []
            finally:
                pages_text_list.append("".join(curr_lines))

        if len(pages_text_list) > 0:
            page_text_o = "".join(pages_text_list)
            self.last_extracted_text = page_text_o
            return page_text_o 
        else:
            raise DocumentException(f"Error: Unable to retrieve text from the section '{section_frame_i}'")
    

    ##########################################################
    # Private methods
    ##########################################################

    def __extract_heading_level(self, heading_i: str) -> int:

        level = 1
        match = re.search(r'\d+', heading_i)
        if match:
            level = int(match.group())
        return level

    ##################################################

    def __build_toc(self) -> list:
        paragraphs_data = []
        total_tokens = 0
        local_toc = []
        
        # Extract all paragraphs and track tokens to map headings to pages
        for para in self.doc.paragraphs:
            text = para.text
            if not text.strip():
                continue
                
            tokens = re.findall(r'\S+\s*', text + '\n')
            paragraphs_data.append({
                'text': text,
                'tokens': tokens,
                'style': para.style.name if para.style else '',
                'token_start': total_tokens
            })
            total_tokens += len(tokens)
            
        # Calculate tokens per page based on the total word count and metadata pages
        tokens_per_page = math.ceil(total_tokens / self.page_num) if self.page_num > 0 else 1
        if tokens_per_page == 0:
            tokens_per_page = 1

        # Build TOC from headings
        for p in paragraphs_data:
            style_name = p['style']
            if style_name and style_name.startswith('Heading'):
                level = self.__extract_heading_level(style_name)
                    
                page_idx = p['token_start'] // tokens_per_page
                local_toc.append([level, p['text'].strip(), page_idx + 1])
                
        # Append an EOF entry to properly close the last section
        local_toc.append([1, "EOF", self.page_num + 1])

        return local_toc

####################################################################################################

if __name__ == '__main__':
    logging.basicConfig(level=logging.INFO, format="[%(name)s :  %(levelname)s] %(message)s")
    
    try:
        # We can test with a dummy docx file if it exists, or create one.        
        wd = WordDocument('/Users/calogeroforte/Local_database/Dante.docx')
        logger.info(f"Loaded '{wd.file_name}' with {wd.page_num} pages.")
        
        if len(wd.toc) > 1:
            logger.info("TOC parsed successfully!")

            print(wd.formtat_toc(wd.toc))

            
            ch1 = wd.get_section_text_by_heading('Canto IX')
            logger.info(f"Chapter IX extracted: {len(ch1[1]) if isinstance(ch1, tuple) else len(ch1)} chars.")
            wd.save_last_extracted_text('/Users/calogeroforte/Local_database/Dante_ch1.txt')
            
        
    except Exception as e:
        logger.error(f"Testing failed: {e}")
    finally:
        if os.path.exists('test_doc.docx'):
            os.remove('test_doc.docx')
